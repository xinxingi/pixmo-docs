import subprocess
import tempfile
import os
import random
from pdf2image import convert_from_bytes
from shutil import rmtree
from PIL import ImageOps
from docx2pdf import convert


def crop_whitespace(image):
    # Invert the image (assuming white background)
    inverted_image = ImageOps.invert(image.convert("RGB"))

    # Convert image to grayscale
    grayscale_image = inverted_image.convert("L")

    # Get bounding box of non-white pixels
    bbox = grayscale_image.getbbox()

    if bbox:
        # Expand bounding box with a 20px buffer
        buffer_size = random.randint(10, 50)
        left = max(0, bbox[0] - buffer_size)
        upper = max(0, bbox[1] - buffer_size)
        right = min(image.width, bbox[2] + buffer_size)
        lower = min(image.height, bbox[3] + buffer_size)

        # Crop image using calculated bounding box
        cropped_image = image.crop((left, upper, right, lower))
        return cropped_image

    # If no bounding box found (all white image), return original image
    return image


def render_docx(docx_object):
    # Create temporary directory as current directory
    temp_dir = "./"

    # Prepare paths
    docx_file = os.path.join(temp_dir, "temp.docx")
    pdf_file = os.path.join(temp_dir, "temp.pdf")

    # Save the docx object to a file
    docx_object.save(docx_file)

    # Convert DOCX to PDF
    convert(docx_file, pdf_file)

    # Convert PDF to image
    images = convert_from_bytes(open(pdf_file, "rb").read())

    # Cleanup temporary files
    os.remove(docx_file)
    os.remove(pdf_file)
    rmtree(temp_dir, ignore_errors=True)

    # Return the PIL image (assuming single-page document)
    if images:
        return crop_whitespace(images[0])  # Return the first page as a PIL image

    raise RuntimeError("DOCX did not generate anything.")


if __name__ == "__main__":
    # Example usage
    from docx import Document

    doc = Document()
    doc.add_heading("Hello, World!", level=1)
    doc.add_paragraph("This is a sample paragraph.")
    image = render_docx(doc)
    image.show()